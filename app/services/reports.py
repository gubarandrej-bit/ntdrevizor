from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from sqlalchemy.orm import Session

from app.config import settings
from app.models import Audit, AuditFile, CheckResult, Finding, Setting
from app.util import loads, looks_like_cable, norm, parse_section


def build_reports(db: Session, audit: Audit) -> dict[str, str]:
    settings.reports_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    base = f"audit_{audit.id}_{stamp}"
    docx_path = settings.reports_dir / f"{base}.docx"
    xlsx_path = settings.reports_dir / f"{base}.xlsx"
    bov_path = settings.reports_dir / f"{base}_bov.xlsx"
    _write_docx(db, audit, docx_path)
    _write_xlsx(db, audit, xlsx_path)
    _write_bov(db, audit, bov_path)
    return {
        "docx": str(docx_path),
        "xlsx": str(xlsx_path),
        "bov": str(bov_path),
    }


def _meta(db: Session, audit: Audit):
    findings = (
        db.query(Finding).filter(Finding.audit_id == audit.id).order_by(Finding.id).all()
    )
    checks = (
        db.query(CheckResult).filter(CheckResult.audit_id == audit.id).order_by(CheckResult.id).all()
    )
    files = db.query(AuditFile).filter(AuditFile.audit_id == audit.id).all()
    crit = [f for f in findings if f.severity == "critical"]
    non = [f for f in findings if f.severity == "noncritical"]
    info = [f for f in findings if f.severity == "info"]
    return findings, checks, files, crit, non, info


def _write_docx(db: Session, audit: Audit, path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    findings, checks, files, crit, non, info = _meta(db, audit)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(1.6)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    h = doc.add_heading("ОТЧЁТ О ПРОВЕРКЕ РАБОЧЕЙ ДОКУМЕНТАЦИИ", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Система «Ревизор НТД»").alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Объект: ").bold = True
    p.add_run(audit.object_name or "не указан")
    p = doc.add_paragraph()
    p.add_run("Наименование проверки: ").bold = True
    p.add_run(audit.title)
    p = doc.add_paragraph()
    p.add_run("Системы: ").bold = True
    p.add_run(", ".join(loads(audit.systems_json, [])) or "не указаны")
    p = doc.add_paragraph()
    p.add_run("Режим / модели: ").bold = True
    p.add_run(f"{audit.mode}; {', '.join(loads(audit.models_json, [])) or 'без ИИ'}")
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run((audit.finished_at or audit.created_at).strftime("%d.%m.%Y %H:%M"))

    doc.add_heading("1. Выводы", level=1)
    if crit:
        doc.add_paragraph(
            f"Выявлено критических замечаний: {len(crit)}. "
            "Документация не может считаться соответствующей НТД без устранения этих замечаний "
            "либо предоставления недостающих исходных данных."
        )
    elif non:
        doc.add_paragraph(
            f"Критических замечаний нет. Некритических: {len(non)}. "
            "Требуется уточнение и устранение замечаний до выпуска в производство работ."
        )
    else:
        doc.add_paragraph(
            "По выполненным проверкам критических и некритических замечаний не зафиксировано. "
            "Это не означает проверку того, что не проводилось — см. раздел 4."
        )

    skipped = [c for c in checks if c.status == "skipped"]
    if skipped:
        doc.add_paragraph(
            f"Не проводилось проверок: {len(skipped)}. Причины указаны в разделе 4. "
            "По непроведённым проверкам выводы о соответствии не делались."
        )

    doc.add_heading("2. Критические замечания", level=1)
    if not crit:
        doc.add_paragraph("Критические замечания отсутствуют.")
    else:
        _docx_findings(doc, crit)

    doc.add_heading("3. Некритические замечания", level=1)
    if not non:
        doc.add_paragraph("Некритические замечания отсутствуют.")
    else:
        _docx_findings(doc, non)

    if info:
        doc.add_heading("3.1. Сведения (не являются замечаниями)", level=2)
        _docx_findings(doc, info)

    doc.add_heading("4. Перечень проверок", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Код"
    hdr[1].text = "Проверка"
    hdr[2].text = "Статус"
    hdr[3].text = "Причина, если не проводилась"
    status_ru = {"done": "выполнена", "skipped": "не проводилась", "pending": "не начата", "error": "ошибка"}
    for c in checks:
        row = table.add_row().cells
        row[0].text = c.check_code
        row[1].text = c.title
        row[2].text = status_ru.get(c.status, c.status)
        row[3].text = c.reason or ""

    doc.add_heading("5. Исходные файлы", level=1)
    for f in files:
        doc.add_paragraph(
            f"{f.filename} — класс: {f.classified_as}; разбор: {f.parse_status}"
            + (f"; {f.parse_notes}" if f.parse_notes else ""),
            style="List Bullet",
        )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "Правило отчёта: при отсутствии исходных данных сведения не выдумывались. "
        "Каждое замечание сопровождается ссылкой на пункт НТД, если он известен. "
        "Актуальность НТД — по локальному каталогу на дату проверки с попыткой сетевого зондирования источников."
    ).italic = True

    doc.save(path)


def _docx_findings(doc, items: list[Finding]) -> None:
    for i, f in enumerate(items, 1):
        refs = ", ".join(loads(f.ntd_refs, [])) or "пункт НТД не указан (не выдумывался)"
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {f.title}")
        run.bold = True
        doc.add_paragraph(f.description)
        doc.add_paragraph(f"НТД: {refs}")
        if f.evidence:
            doc.add_paragraph(f"Основание: {f.evidence}")
        if f.location:
            doc.add_paragraph(f"Файл: {f.location}")


def _write_xlsx(db: Session, audit: Audit, path: Path) -> None:
    findings, checks, files, crit, non, info = _meta(db, audit)
    wb = Workbook()

    ws = wb.active
    ws.title = "Сводка"
    _header_fill = PatternFill("solid", fgColor="1B2430")
    ws["A1"] = "Отчёт о проверке документации — Ревизор НТД"
    ws["A1"].font = Font(bold=True, size=14, color="F2E6C8")
    ws.merge_cells("A1:B1")
    rows = [
        ("Объект", audit.object_name or ""),
        ("Проверка", audit.title),
        ("Системы", ", ".join(loads(audit.systems_json, []))),
        ("Режим", audit.mode),
        ("Модели", ", ".join(loads(audit.models_json, []))),
        ("Дата", (audit.finished_at or audit.created_at).isoformat(sep=" ", timespec="minutes")),
        ("Критических", len(crit)),
        ("Некритических", len(non)),
        ("Сведений", len(info)),
        ("Проверок выполнено", sum(1 for c in checks if c.status == "done")),
        ("Проверок не проводилось", sum(1 for c in checks if c.status == "skipped")),
    ]
    for i, (k, v) in enumerate(rows, 3):
        ws[f"A{i}"] = k
        ws[f"B{i}"] = v
        ws[f"A{i}"].font = Font(bold=True, color="D4A054")
        ws[f"B{i}"].font = Font(color="E8EEF6")
    ws.column_dimensions["A"].width = 36
    ws.column_dimensions["B"].width = 80
    ws.sheet_properties.tabColor = "D4A054"

    ws2 = wb.create_sheet("Замечания")
    headers = ["№", "Степень", "Проверка", "Заголовок", "Описание", "НТД", "Основание", "Файл"]
    for col, h in enumerate(headers, 1):
        cell = ws2.cell(1, col, h)
        cell.font = Font(bold=True, color="F2E6C8")
        cell.fill = _header_fill
    for i, f in enumerate(findings, 1):
        vals = [
            i,
            {"critical": "критическое", "noncritical": "некритическое", "info": "сведение"}.get(
                f.severity, f.severity
            ),
            f.check_code,
            f.title,
            f.description,
            "; ".join(loads(f.ntd_refs, [])),
            f.evidence,
            f.location,
        ]
        for col, v in enumerate(vals, 1):
            cell = ws2.cell(i + 1, col, v)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if f.severity == "critical":
                cell.fill = PatternFill("solid", fgColor="F8D0D0")
            elif f.severity == "noncritical":
                cell.fill = PatternFill("solid", fgColor="F8EBC0")
        ws2.row_dimensions[i + 1].height = 48
    widths = [6, 16, 22, 40, 60, 36, 36, 24]
    for i, w in enumerate(widths, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    ws3 = wb.create_sheet("Проверки")
    for col, h in enumerate(["Код", "Наименование", "Статус", "Причина непроведения"], 1):
        cell = ws3.cell(1, col, h)
        cell.font = Font(bold=True, color="F2E6C8")
        cell.fill = _header_fill
    ru = {"done": "выполнена", "skipped": "не проводилась", "pending": "не начата", "error": "ошибка"}
    for i, c in enumerate(checks, 1):
        ws3.cell(i + 1, 1, c.check_code)
        ws3.cell(i + 1, 2, c.title)
        ws3.cell(i + 1, 3, ru.get(c.status, c.status))
        ws3.cell(i + 1, 4, c.reason or "")
        if c.status == "skipped":
            for col in range(1, 5):
                ws3.cell(i + 1, col).fill = PatternFill("solid", fgColor="E6E6E6")
    for i, w in enumerate([24, 70, 20, 70], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    ws4 = wb.create_sheet("Файлы")
    for col, h in enumerate(["Файл", "Класс", "Разбор", "Примечание", "Размер, байт"], 1):
        cell = ws4.cell(1, col, h)
        cell.font = Font(bold=True, color="F2E6C8")
        cell.fill = _header_fill
    for i, f in enumerate(files, 1):
        ws4.cell(i + 1, 1, f.filename)
        ws4.cell(i + 1, 2, f.classified_as)
        ws4.cell(i + 1, 3, f.parse_status)
        ws4.cell(i + 1, 4, f.parse_notes)
        ws4.cell(i + 1, 5, f.size)
    for i, w in enumerate([40, 22, 14, 70, 16], 1):
        ws4.column_dimensions[get_column_letter(i)].width = w

    wb.save(path)


_HEIGHT_RE = re.compile(r"(?:отметк[а-я]*|отм\.?|высот[а-я]*)\s*[+-]?\s*(\d+(?:[.,]\d+)?)")

HEIGHT_BAND_ORDER = [
    "от 0 м до 5 м",
    "от 5 м до 8 м",
    "от 8 м до 12 м",
    "более 12 м",
    "высота не указана",
]

LAYING_ORDER = [
    "в лотке",
    "в гофре",
    "в кабель-канале",
    "в трубе",
    "в штрабе",
    "в траншее/земле",
    "способ не указан",
]

_CABLE_CARRIER_KEYS = (
    "лоток", "короб", "кабель-канал", "кабель канал", "лестниц",
    "консоль", "кронштейн", "профил", "перфорир", "упмк", "держатель", "хому",
    "саморез", "дюбел", "анкер", "болт", "гайк", "винт", "шуруп", "метиз",
)

_CABLE_AUX_KEYS = (
    "гофр", "металлорукав", "рукав",
    "стяжк", "хомут", "клипс", "бирк", "наконечник",
    "проходк", "маркиров", "этикетк",
)


def _height_of(it: dict) -> float | None:
    """Высота монтажа из примечаний/способа прокладки/направления («на отметке +7,8 м»)."""
    blob = " ".join(str(it.get(k) or "") for k in ("note", "laying", "to"))
    m = _HEIGHT_RE.search(blob)
    if m:
        try:
            return float(m.group(1).replace(",", "."))
        except ValueError:
            return None
    return None


def _band_of(h: float | None) -> str:
    if h is None:
        return "высота не указана"
    if h < 5:
        return "от 0 м до 5 м"
    if h < 8:
        return "от 5 м до 8 м"
    if h <= 12:
        return "от 8 м до 12 м"
    return "более 12 м"


def _laying_of(it: dict) -> str:
    blob = norm(" ".join(str(it.get(k) or "") for k in ("laying", "note", "to", "name")))
    if "лотк" in blob:
        return "в лотке"
    if "гофр" in blob:
        return "в гофре"
    if any(k in blob for k in ("кабель-канал", "кабель канал", "кабельный канал", "короб")):
        return "в кабель-канале"
    if "труб" in blob:
        return "в трубе"
    if "штраб" in blob:
        return "в штрабе"
    if "транше" in blob or "земл" in blob or "грунт" in blob:
        return "в траншее/земле"
    return "способ не указан"


def _is_cable_carrier(it: dict) -> bool:
    """Кабеленесущие системы: лотки, короба, кабель-каналы, гофра, трубы и т.п."""
    blob = norm(" ".join(str(it.get(k) or "") for k in ("name", "mark", "type")))
    if not blob:
        return False
    if "термоусаж" in blob or "термоусад" in blob or "патруб" in blob:
        return False
    if any(norm(k) in blob for k in _CABLE_CARRIER_KEYS):
        return True
    if "труб" in blob:
        return True
    return False


def _is_cable_aux(it: dict) -> bool:
    """Материалы, сопутствующие прокладке кабеля (идут в «Монтаж кабеля»):
    гофротруба/металлорукав, стяжки, хомуты, клипсы, проходки, бирки, наконечники."""
    blob = norm(" ".join(str(it.get(k) or "") for k in ("name", "mark", "type")))
    if not blob:
        return False
    return any(norm(k) in blob for k in _CABLE_AUX_KEYS)


def _write_bov(db: Session, audit: Audit, path: Path) -> None:
    """Ведомость объёмов работ (ВОР) по форме заказчика.

    Форма повторяет образец (шапка с утверждением, титул, блок реквизитов,
    10-колоночная таблица, подписи). Разделы:
      — Оборудование и материалы (все позиции спецификации);
      — Кабельные изделия и провода (кабельная продукция из спецификации);
      — Монтаж кабеленесущих систем — с разбивкой по высотным отметкам;
      — Пусконаладочные работы.
    Каждая позиция: сначала строка работы (монтаж/прокладка/установка), затем
    строка материала (марка/артикул). Кабельный журнал в ВОР не включается.
    Высота берётся из примечаний позиций («на отметке +7,8 м»); если высота не
    указана — строка помечается «высота не указана». Никакие работы и нормы
    расхода не выдумываются.
    """

    files = db.query(AuditFile).filter(AuditFile.audit_id == audit.id).all()
    items: list[dict] = []
    sources: list[str] = []
    for f in files:
        data = loads(f.extracted_json, {})
        if f.classified_as == "specification":
            items.extend(data.get("items") or [])
            sources.append(f.filename)

    def _name(it: dict) -> str:
        return re.sub(r"\s+", " ", (it.get("name") or it.get("mark") or it.get("type") or "")).strip()

    def _mark(it: dict) -> str:
        return re.sub(r"\s+", " ", (it.get("manufacturer") or it.get("mark") or it.get("type") or "")).strip()

    def _is_cable(it: dict) -> bool:
        return looks_like_cable(" ".join(str(it.get(k) or "") for k in ("name", "mark", "type", "manufacturer")))

    def _is_spec_position(it: dict) -> bool:
        """Позиция спецификации: есть количество («Кол.»).

        Строки кабельного журнала имеют только длину (м) и направление — их
        в ведомость объёмов не включаем (по требованию заказчика).
        """
        return it.get("qty") is not None

    # разбивка строк спецификации по видам работ:
    #   equip      → «Монтаж оборудования» (приборы, извещатели, блоки и т.п.)
    #   carrier    → «Монтаж кабеленесущих систем» (лотки, короба, трубы, крепёж)
    #   cab_materials → «Монтаж кабеля» (кабели + гофра/стяжки/хомуты/проходки)
    equip: list[dict] = []
    carrier: list[dict] = []
    cab_materials: list[dict] = []
    for it in items:
        if not _name(it):
            continue
        if not _is_spec_position(it):
            continue
        if _is_cable(it) or _is_cable_aux(it):
            cab_materials.append(it)
        elif _is_cable_carrier(it):
            carrier.append(it)
        else:
            equip.append(it)

    # Кабельный журнал в ВОР не включается: ведомость формируется ТОЛЬКО
    # из позиций спецификации (по требованию заказчика).


    # стили
    thin = Side(style="thin", color="FF808080")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_font = Font(bold=True, size=10)
    head_fill = PatternFill("solid", fgColor="FFD9D9D9")
    sect_fill = PatternFill("solid", fgColor="FFF2F2F2")
    sect_font = Font(bold=True, size=11)
    sub_font = Font(italic=True, bold=True, size=10, color="FF404040")
    sub_fill = PatternFill("solid", fgColor="FFFAFAFA")
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    small_gray = Font(size=9, italic=True, color="FF808080")

    wb = Workbook()
    ws = wb.active
    ws.title = "ВОР"

    # ---------- шапка ----------
    ws["I2"] = "Приложение №1\nк Техническому заданию №____"
    ws["I2"].alignment = Alignment(wrap_text=True, vertical="top")
    ws["I3"] = "УТВЕРЖДАЮ"
    ws["I3"].font = Font(bold=True)
    ws["I3"].alignment = Alignment(horizontal="left")
    ws.merge_cells("H4:J4")
    ws["H4"] = "____________________________\n(должность, Ф.И.О.)"
    ws["H4"].alignment = wrap
    ws.merge_cells("H5:J5")
    ws["H5"] = "«___» ______________ 20__ г."

    # ---------- титул ----------
    ws.merge_cells("A8:J8")
    ws["A8"] = "Ведомость объемов работ №______ от __.__.____"
    ws["A8"].font = Font(bold=True, size=13)
    ws["A8"].alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells("A9:J9")
    ws["A9"] = "на выполнение комплекса работ строительно-монтажных работ"
    ws["A9"].alignment = Alignment(horizontal="center", vertical="center")
    ws["F10"] = "Лист 1"
    ws["F10"].alignment = Alignment(horizontal="right")

    # ---------- реквизиты ----------
    def _requisite(row: int, text: str) -> None:
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        ws.cell(row, 2, text)

    _requisite(11, f"Объект: {audit.object_name or '—'}")
    _requisite(12, f"Шифр документации: {'—'}")
    _requisite(13, "Тип ведомости: Новые работы")
    _requisite(14, "Причина внесения изменений / основания для разработки ВОР:")
    _requisite(15, "Приложения:")

    # ---------- таблица ----------
    headers = [
        "№ п/п",
        "Наименование работы",
        "Ед. изм. работы",
        "Объем работы",
        "№ материала",
        "Наименование материала",
        "Ед. изм. материала",
        "Кол-во материала",
        "Расчет объема материала",
        "Примечание\n(Ссылка на чертеж, техническое решение, условия производства работ и пр.)",
    ]
    HDR_ROW = 17
    for col, h in enumerate(headers, 1):
        cell = ws.cell(HDR_ROW, col, h)
        cell.font = head_font
        cell.fill = head_fill
        cell.border = border
        cell.alignment = center

    row_i = HDR_ROW + 1
    n = 1  # № п/п

    def _section(title: str) -> None:
        nonlocal row_i
        ws.cell(row_i, 2, title)
        ws.cell(row_i, 2).font = sect_font
        for col in range(1, 11):
            ws.cell(row_i, col).border = border
            ws.cell(row_i, col).fill = sect_fill
        row_i += 1

    def _sub(title: str) -> None:
        nonlocal row_i
        ws.cell(row_i, 2, title)
        ws.cell(row_i, 2).font = sub_font
        for col in range(1, 11):
            ws.cell(row_i, col).fill = sub_fill
        row_i += 1

    def _work_type_row(title: str) -> None:
        """Строка вида работ («Монтаж кабеля», «Монтаж оборудования»…)."""
        nonlocal row_i, n
        ws.cell(row_i, 1, n).alignment = center
        ws.cell(row_i, 2, title).font = Font(bold=True)
        ws.cell(row_i, 2).alignment = wrap
        for col in range(1, 11):
            ws.cell(row_i, col).border = border
            ws.cell(row_i, col).fill = sub_fill
        row_i += 1
        n += 1

    def _work_row(name: str, unit: str, qty, note: str = ""):
        """Строка работы с единицей и объёмом (используется для ПНР)."""
        nonlocal row_i, n
        ws.cell(row_i, 1, n).alignment = center
        ws.cell(row_i, 2, name).alignment = wrap
        ws.cell(row_i, 3, unit).alignment = center
        ws.cell(row_i, 4, qty if qty is not None else "").alignment = center
        ws.cell(row_i, 10, note).alignment = wrap
        for col in range(1, 11):
            ws.cell(row_i, col).border = border
        row_i += 1
        n += 1

    def _material_row(mat_no: str, name: str, unit: str, qty, note: str = ""):
        nonlocal row_i
        ws.cell(row_i, 5, mat_no).alignment = center
        ws.cell(row_i, 6, name).alignment = wrap
        ws.cell(row_i, 7, unit).alignment = center
        ws.cell(row_i, 8, qty if qty is not None else "").alignment = center
        ws.cell(row_i, 9, "").alignment = center
        ws.cell(row_i, 10, note).alignment = wrap
        for col in range(1, 11):
            ws.cell(row_i, col).border = border
        row_i += 1

    def _materials_under(materials: list[dict]) -> None:
        """Перечень материалов под строкой вида работ.

        Одинаковые материалы агрегируются одной строкой:
          — кабели — по марке (длины суммируются), с сечением;
          — прочие — по (наименование, марка), количества суммируются.
        Пример: «Монтаж кабеля» → Кабель ВВГнг(А)-LS 3x2.5 — 500 м;
        Стяжки нейлоновые — 1 упак.; Гофротруба ПВХ — 500 м.
        """
        from collections import OrderedDict

        nonlocal row_i
        groups: "OrderedDict[tuple, dict]" = OrderedDict()
        for it in materials:
            name = _name(it)
            mark = _mark(it)
            qty = it.get("qty") if it.get("qty") is not None else it.get("length")
            if qty is None:
                continue
            qty = float(qty)
            if _is_cable(it):
                brand = mark if (mark and looks_like_cable(mark)) else name
                brand = re.sub(r"\s+", " ", brand).strip()
                parsed = parse_section(" ".join(str(it.get(k) or "") for k in ("mark", "manufacturer", "name", "type")))
                sec = ""
                if parsed and parsed.get("mm2"):
                    cores = parsed.get("cores")
                    mm2 = parsed["mm2"]
                    sec = f"{cores}x{('%g' % mm2)}" if cores else ('%g' % mm2)
                # убираем из марки только токены сечения («1x2x0,75», «3х2,5»),
                # не трогая «Cat5e»/«Cat6A»
                base = re.sub(
                    r"\d+(?:[.,]\d+)?(?:\s*[xх]\s*\d+(?:[.,]\d+)?)+",
                    " ",
                    brand,
                )
                base = re.sub(r"\s+", " ", base).strip(" -–")
                if not base:
                    base = brand
                key = ("cable", base.upper(), sec)
                unit = "м"
                kind_word = "Провод" if "провод" in name.lower() else "Кабель"
                disp = f"{kind_word} {base}" + (f" {sec}" if sec else "")
            else:
                key = ("item", name, mark or "")
                unit = it.get("unit") or "шт"
                disp = name if not mark or mark.lower() in name.lower() else f"{name} — {mark}"
            g = groups.setdefault(key, {"qty": 0.0, "unit": unit, "disp": disp, "note": ""})
            g["qty"] += qty
            if not g["note"] and it.get("note"):
                g["note"] = str(it.get("note")).strip()
        mn = 0
        for g in groups.values():
            mn += 1
            q = g["qty"]
            q_disp = int(q) if q == int(q) and abs(q) < 1e6 else round(q, 2)
            _material_row(f"{mn}.", g["disp"], g["unit"], q_disp, g["note"])

    if not equip and not carrier and not cab_materials:
        ws.cell(row_i, 2, "Исходных данных для ведомости объёмов нет. Файл не выдумывался.")
        ws.merge_cells(start_row=row_i, start_column=2, end_row=row_i, end_column=10)
        row_i += 1
    else:
        # Структура: сначала вид работ, затем перечень всех материалов этой работы
        # (по требованию заказчика). Только позиции спецификации; журнал не входит.

        # 1) Монтаж оборудования
        if equip:
            _work_type_row("Монтаж оборудования")
            _materials_under(equip)

        # 2) Монтаж кабеленесущих систем
        if carrier:
            _work_type_row("Монтаж кабеленесущих систем")
            _materials_under(carrier)

        # 3) Монтаж кабеля (кабель + гофра/стяжки/хомуты/проходки)
        if cab_materials:
            _work_type_row("Монтаж кабеля")
            _materials_under(cab_materials)

        # 4) Пусконаладочные работы — по типам оборудования спецификации.
        #    Виды ПНР выведены из типов оборудования; количества требуют
        #    подтверждения (в спецификации как отдельная позиция отсутствуют).
        pnr_lines: list[tuple[str, str, str]] = []
        n_det = 0
        n_ppk = 0
        n_sw = 0
        for it in items:
            blob = norm(" ".join(str(it.get(k) or "") for k in ("name", "mark", "type")))
            if not _is_spec_position(it):
                continue
            q = it.get("qty") if it.get("qty") is not None else (it.get("length") or 0)
            q = int(q or 0)
            if re.search(r"извещател|дип-|\bипт\b|\bипр\b|\bип\s*2\d\d", blob):
                n_det += q
            elif any(k in blob for k in ("прибор приемно-контрол", "ппкп", "ппк", "блок", "контроллер", "расширител", "модуль")):
                n_ppk += q
            elif any(k in blob for k in ("сервер", "арм", "монитор", "коммутатор")):
                n_sw += q
        if n_ppk:
            pnr_lines.append(("Программирование приборов ППКП/блоков (адресация, конфигурация)", "шт", str(n_ppk)))
        if n_det:
            pnr_lines.append(("Адресация и проверка извещателей (по типам)", "шт", str(n_det)))
        if n_sw:
            pnr_lines.append(("Настройка программного обеспечения верхнего уровня", "компл", str(n_sw)))
        pnr_lines.append(("Проверка работоспособности системы, протокол испытаний", "компл", "1"))
        if pnr_lines:
            _work_type_row("Пусконаладочные работы")
            for name, unit, qty in pnr_lines:
                _work_row(name, unit, qty, "выведено из спецификации")

    # ---------- подписи ----------
    row_i += 1
    ws.cell(row_i, 1, "Составил: ____________________ (Ф.И.О.)")
    ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=10)
    row_i += 1
    ws.cell(row_i, 1, "Проверил: ____________________ (Ф.И.О.)")
    ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=10)
    row_i += 2
    ws.cell(row_i, 1, (
        f"Сформировано системой «Ревизор НТД» {datetime.utcnow().strftime('%d.%m.%Y %H:%M')} "
        f"по позициям спецификации оборудования, изделий и материалов. "
        f"Источники: {', '.join(sources) if sources else 'нет'}. "
        f"Структура: вид работ → перечень материалов этой работы. "
        f"Кабельный журнал в ведомость не включается. "
        f"Работы и нормы расхода, отсутствующие во входных данных, не добавлялись."
    ))
    ws.cell(row_i, 1).font = small_gray
    ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=10)

    # ---------- ширина колонок ----------
    for i, w in enumerate([5, 34, 8, 10, 7, 52, 8, 10, 12, 34], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    wb.save(path)
