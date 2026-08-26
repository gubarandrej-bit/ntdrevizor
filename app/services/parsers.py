from __future__ import annotations

import json
import logging
import math
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from pathlib import Path
from typing import Any

from app.util import (
    looks_like_cable,
    parse_float,
    parse_section,
    truncate,
)

MAX_TEXT = 200_000
MAX_ROWS = 8_000


def _quiet_pypdf() -> None:
    """Глушим предупреждения pypdf о битых числах в PDF.

    В некоторых PDF (экспорт из CAD) координаты текста склеиваются в строки вида
    «779.045628.445» (две десятичные точки). pypdf не может преобразовать их в
    число и пишет «could not convert string to float … use 0.0 instead» в лог.
    Это не ошибка извлечения: текст читается, координаты подставляются нулевыми.
    Чтобы не засорять журнал сервера сотнями строк — уровень логгера pypdf
    поднимаем до ERROR.
    """
    try:
        logging.getLogger("pypdf").setLevel(logging.ERROR)
    except Exception:
        pass


# Воркер для извлечения текста PDF в отдельном процессе. Некоторые страницы
# (планы с тысячами векторных объектов) разбираются pypdf десятки секунд;
# в потоке их невозможно прервать (daemon-потоки продолжают работать и давят
# на GIL, выжигая бюджет разбора таблиц). Отдельный процесс можно убить целиком.
_PYPDF_TEXT_WORKER = r"""
import json, sys, threading
from pypdf import PdfReader

def main():
    path, out, page_timeout = sys.argv[1], sys.argv[2], float(sys.argv[3])
    reader = PdfReader(path)
    f = open(out, "w", encoding="utf-8")
    for i, page in enumerate(reader.pages[:80]):
        box = {"done": False, "text": ""}
        def job():
            try:
                box["text"] = page.extract_text() or ""
            except Exception:
                box["text"] = ""
            finally:
                box["done"] = True
        t = threading.Thread(target=job, daemon=True)
        t.start()
        t.join(page_timeout)
        if box["done"]:
            if box["text"].strip():
                f.write(json.dumps([i + 1, box["text"]]) + "\n")
                f.flush()
        # «медленные» страницы не записываем — процесс всё равно убьют по бюджету
    f.close()

if __name__ == "__main__":
    main()
"""


def _extract_pdf_text_subprocess(path: Path, budget: float, page_timeout: float) -> tuple[list[str], list[str]]:
    """Извлекает текст страниц PDF в отдельном процессе с жёстким бюджетом.

    Возвращает (text_parts, notes), где text_parts — строки «--- страница N ---».
    Процесс принудительно убивается по бюджету; утечки потоков внутри процесса
    не влияют на основной процесс и последующий разбор таблиц.
    """
    import os

    tmp = Path(tempfile.mkdtemp(prefix="pdftext_"))
    worker = tmp / "worker.py"
    out = tmp / "out.jsonl"
    notes: list[str] = []
    try:
        worker.write_text(_PYPDF_TEXT_WORKER, encoding="utf-8")
        proc = subprocess.Popen(
            [str(sys.executable), str(worker), str(path), str(out), f"{page_timeout}"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        try:
            proc.wait(timeout=budget)
        except subprocess.TimeoutExpired:
            proc.kill()
            try:
                proc.wait(timeout=5)
            except Exception:
                pass
            notes.append(f"Бюджет {budget:g}с на текст PDF исчерпан — часть страниц пропущена.")
        parts: list[str] = []
        if out.exists():
            for line in out.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if not line:
                    continue
                try:
                    pn, text = json.loads(line)
                    if str(text).strip():
                        parts.append(f"--- страница {pn} ---\n{text}")
                except Exception:
                    continue
        return parts, notes
    except Exception as exc:
        notes.append(f"pdf-text-subprocess: {exc}")
        return [], notes
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _text_is_garbled(text: str) -> bool:
    """Определяет «битый» посимвольный текст (экспорт из CAD/XPS).

    В таких PDF каждый символ — отдельный текстовый объект, и extract_text
    возвращает «И\\nн\\nв\\n.\\n...» (почти каждая строка длиной 1–2 символа).
    Для нормального текста доля коротких строк низкая.
    """
    if not text:
        return True
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if len(lines) < 40:
        return False
    short = sum(1 for ln in lines if len(ln.strip()) <= 2)
    return short / len(lines) > 0.6


_PDFPLUMBER_TEXT_WORKER = r"""
import json, sys, threading
import pdfplumber

def main():
    path, out, page_timeout = sys.argv[1], sys.argv[2], float(sys.argv[3])
    f = open(out, "w", encoding="utf-8")
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages[:80]):
            box = {"done": False, "text": ""}
            def job():
                try:
                    box["text"] = page.extract_text() or ""
                except Exception:
                    box["text"] = ""
                finally:
                    box["done"] = True
            t = threading.Thread(target=job, daemon=True)
            t.start()
            t.join(page_timeout)
            if box["done"] and box["text"].strip():
                f.write(json.dumps([i + 1, box["text"]]) + "\n")
                f.flush()
    f.close()

if __name__ == "__main__":
    main()
"""


_PYMUPDF_TEXT_WORKER = r"""
import json, sys
import pymupdf

def main():
    path, out = sys.argv[1], sys.argv[2]
    doc = pymupdf.open(path)
    f = open(out, "w", encoding="utf-8")
    for i, page in enumerate(doc):
        if i >= 80:
            break
        try:
            t = page.get_text()
        except Exception:
            t = ""
        if t.strip():
            f.write(json.dumps([i + 1, t]) + "\n")
            f.flush()
    doc.close()
    f.close()

if __name__ == "__main__":
    main()
"""


def _extract_pymupdf_text_subprocess(path: Path, budget: float) -> tuple[list[str] | None, list[str]]:
    """Извлекает текст через PyMuPDF — быстро и корректно для PDF из CAD/XPS.

    Возвращает (None, notes), если пакет не установлен."""
    try:
        import pymupdf  # noqa: F401
    except ImportError:
        return None, ["PyMuPDF не установлен (pip install pymupdf)."]
    tmp = Path(tempfile.mkdtemp(prefix="pymupdf_"))
    worker = tmp / "worker.py"
    out = tmp / "out.jsonl"
    notes: list[str] = []
    try:
        worker.write_text(_PYMUPDF_TEXT_WORKER, encoding="utf-8")
        proc = subprocess.Popen(
            [str(sys.executable), str(worker), str(path), str(out)],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        try:
            proc.wait(timeout=budget)
        except subprocess.TimeoutExpired:
            proc.kill()
            try:
                proc.wait(timeout=5)
            except Exception:
                pass
            notes.append(f"Бюджет {budget:g}с на PyMuPDF исчерпан.")
        parts: list[str] = []
        if out.exists():
            for line in out.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if not line:
                    continue
                try:
                    pn, text = json.loads(line)
                    if str(text).strip():
                        parts.append(f"--- страница {pn} ---\n{text}")
                except Exception:
                    continue
        return parts, notes
    except Exception as exc:
        notes.append(f"pymupdf-text-subprocess: {exc}")
        return None, notes
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _extract_text_pdftotext(path: Path, budget: float) -> tuple[list[str] | None, list[str]]:
    """Извлекает текст через pdftotext (poppler) — самый быстрый и корректный
    способ для PDF, экспортированных из CAD/XPS. Возвращает (None, notes), если
    утилита не установлена."""
    exe = shutil.which("pdftotext")
    if not exe:
        return None, ["pdftotext не установлен (пакет poppler-utils)."]
    tmp = Path(tempfile.mkdtemp(prefix="pdftotext_"))
    out = tmp / "out.txt"
    notes: list[str] = []
    try:
        proc = subprocess.Popen(
            [exe, "-layout", str(path), str(out)],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        try:
            proc.wait(timeout=budget)
        except subprocess.TimeoutExpired:
            proc.kill()
            try:
                proc.wait(timeout=5)
            except Exception:
                pass
            notes.append(f"Бюджет {budget:g}с на pdftotext исчерпан.")
        parts: list[str] = []
        if out.exists():
            text = out.read_text(encoding="utf-8", errors="replace")
            for i, pg in enumerate(text.split("\f")):
                if pg.strip():
                    parts.append(f"--- страница {i + 1} ---\n{pg}")
        return parts, notes
    except Exception as exc:
        notes.append(f"pdftotext: {exc}")
        return None, notes
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _extract_pdfplumber_text_subprocess(path: Path, budget: float, page_timeout: float) -> tuple[list[str], list[str]]:
    """Извлекает текст страниц PDF через pdfplumber в отдельном процессе.

    pdfplumber корректно собирает слова в PDF, экспортированных из CAD/XPS
    (где pypdf даёт посимвольный текст). Бюджет и постраничный тайм-аут — как
    у pypdf-варианта.
    """
    tmp = Path(tempfile.mkdtemp(prefix="pdftext_pp_"))
    worker = tmp / "worker.py"
    out = tmp / "out.jsonl"
    notes: list[str] = []
    try:
        worker.write_text(_PDFPLUMBER_TEXT_WORKER, encoding="utf-8")
        proc = subprocess.Popen(
            [str(sys.executable), str(worker), str(path), str(out), f"{page_timeout}"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        try:
            proc.wait(timeout=budget)
        except subprocess.TimeoutExpired:
            proc.kill()
            try:
                proc.wait(timeout=5)
            except Exception:
                pass
            notes.append(f"Бюджет {budget:g}с на текст PDF (pdfplumber) исчерпан — часть страниц пропущена.")
        parts: list[str] = []
        if out.exists():
            for line in out.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if not line:
                    continue
                try:
                    pn, text = json.loads(line)
                    if str(text).strip():
                        parts.append(f"--- страница {pn} ---\n{text}")
                except Exception:
                    continue
        return parts, notes
    except Exception as exc:
        notes.append(f"pdfplumber-text-subprocess: {exc}")
        return [], notes
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def parse_file(path: Path) -> dict[str, Any]:
    ext = path.suffix.lower()
    if ext in {".xls", ".xlsx", ".xlsm"}:
        return parse_spreadsheet(path)
    if ext in {".doc", ".docx"}:
        return parse_document(path)
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".dxf":
        return parse_dxf(path)
    if ext == ".dwg":
        return parse_dwg(path)
    if ext in {".txt", ".csv"}:
        return parse_text(path)
    return {
        "ok": False,
        "kind": "unknown",
        "error": f"Формат {ext or '(без расширения)'} не поддерживается. Допустимы: xls, xlsx, doc, docx, pdf, dwg, dxf.",
        "text": "",
        "tables": [],
        "items": [],
        "cables": [],
        "equipment": [],
        "lengths": [],
        "texts_geom": [],
    }


def parse_text(path: Path) -> dict[str, Any]:
    raw = path.read_bytes()
    text = raw.decode("utf-8", errors="replace")
    if "\x00" in text[:200]:
        text = raw.decode("cp1251", errors="replace")
    return _base(ok=True, kind="text", text=text[:MAX_TEXT])


def parse_spreadsheet(path: Path) -> dict[str, Any]:
    try:
        import openpyxl
    except ImportError as exc:
        return _base(ok=False, kind="xls", error=f"openpyxl недоступен: {exc}")

    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:
        return _base(ok=False, kind="xls", error=f"Не удалось открыть книгу: {exc}")

    tables: list[dict[str, Any]] = []
    text_parts: list[str] = [f"Листы: {', '.join(wb.sheetnames)}"]
    items: list[dict[str, Any]] = []
    cables: list[dict[str, Any]] = []

    try:
        for sheet in wb.worksheets:
            rows = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= MAX_ROWS:
                    break
                values = [_cell(c) for c in row]
                if any(v != "" for v in values):
                    rows.append(values)
            if not rows:
                continue
            header_idx = _header_row(rows)
            headers = rows[header_idx] if header_idx is not None else []
            mapped = _map_headers(headers)
            records = []
            start = (header_idx + 1) if header_idx is not None else 0
            for row in rows[start:]:
                rec = {}
                if mapped:
                    for key, idx in mapped.items():
                        if idx < len(row):
                            rec[key] = row[idx]
                rec["_raw"] = row
                records.append(rec)
            tables.append(
                {
                    "sheet": sheet.title,
                    "headers": headers,
                    "mapped": mapped,
                    "records": records[:MAX_ROWS],
                }
            )
            text_parts.append(f"\n=== Лист {sheet.title} ===")
            for row in rows[:80]:
                text_parts.append(" | ".join(str(x) for x in row if x != ""))

            for rec in records:
                item = _record_to_item(rec, sheet.title)
                if not item:
                    continue
                items.append(item)
                if looks_like_cable(_cable_blob(item)) or rec.get("from") or rec.get("to") or rec.get("length"):
                    cables.append(item)
    finally:
        wb.close()

    return _base(
        ok=True,
        kind="xls",
        text="\n".join(text_parts)[:MAX_TEXT],
        tables=tables,
        items=items,
        cables=cables,
        equipment=[i for i in items if not looks_like_cable(_cable_blob(i))],
    )


def parse_document(path: Path) -> dict[str, Any]:
    ext = path.suffix.lower()
    if ext == ".doc":
        converted = _convert_with_libreoffice(path, "docx")
        if converted is None:
            # last resort: strings
            raw = path.read_bytes()
            text = raw.decode("cp1251", errors="ignore")
            text = "".join(ch if ch.isprintable() or ch in "\n\t" else " " for ch in text)
            text = "\n".join(line.strip() for line in text.splitlines() if len(line.strip()) > 3)
            if len(text) < 40:
                return _base(
                    ok=False,
                    kind="doc",
                    error="Формат .doc (OLE) не разобран: LibreOffice/soffice не найден. Сохраните файл как .docx.",
                )
            return _base(
                ok=True,
                kind="doc",
                text=text[:MAX_TEXT],
                notes="Текст извлечён грубо из .doc без LibreOffice. Таблицы не разобраны.",
            )
        path = converted

    try:
        import docx
    except ImportError as exc:
        return _base(ok=False, kind="docx", error=f"python-docx недоступен: {exc}")

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        return _base(ok=False, kind="docx", error=f"Не удалось открыть документ: {exc}")

    paras = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    tables: list[dict[str, Any]] = []
    items: list[dict[str, Any]] = []
    cables: list[dict[str, Any]] = []
    for ti, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([c.text.strip() for c in row.cells])
        if not rows:
            continue
        header_idx = _header_row(rows)
        headers = rows[header_idx] if header_idx is not None else rows[0]
        mapped = _map_headers(headers)
        records = []
        start = (header_idx + 1) if header_idx is not None else 1
        for row in rows[start:]:
            rec = {}
            for key, idx in mapped.items():
                if idx < len(row):
                    rec[key] = row[idx]
            rec["_raw"] = row
            records.append(rec)
            item = _record_to_item(rec, f"таблица {ti + 1}")
            if item:
                items.append(item)
                if looks_like_cable(_cable_blob(item)):
                    cables.append(item)
        tables.append({"sheet": f"table_{ti + 1}", "headers": headers, "mapped": mapped, "records": records})

    return _base(
        ok=True,
        kind="docx",
        text="\n".join(paras)[:MAX_TEXT],
        tables=tables,
        items=items,
        cables=cables,
        equipment=[i for i in items if not looks_like_cable(_cable_blob(i))],
    )


def _extract_text_bounded(page: Any, timeout: float) -> str | None:
    """extract_text() с тайм-аутом. Возвращает None, если страница «зависла».

    Некоторые страницы с плотной векторной графикой разбираются секунды
    и десятки секунд; их текст пропускаем, чтобы не блокировать проверку.
    """
    box: dict[str, Any] = {"done": False, "text": ""}

    def _job() -> None:
        try:
            box["text"] = page.extract_text() or ""
        except Exception:
            box["text"] = ""
        finally:
            box["done"] = True

    t = threading.Thread(target=_job, daemon=True)
    t.start()
    t.join(timeout)
    return box["text"] if box["done"] else None


_PLAN_PAGE_MARKERS = ("план расположен", "план прокладк", "план трасс")


def _is_cable_color(c: Any) -> bool:
    """Цветная линия (красная/синяя/зелёная и т.п.) — признак трассы на CAD-плане.

    Стены, оси, размеры и штриховки обычно чёрные/серые; кабельные трассы на
    планах СКУД/ЭОМ рисуют цветом. Возвращает False для серых/чёрных/белых.
    """
    if not c:
        return False
    try:
        r, g, b = c
    except Exception:
        return False
    mx, mn = max(c), min(c)
    if mx - mn < 0.15:  # серый/чёрный/белый
        return False
    if mx < 0.3:  # слишком тёмный — почти чёрный
        return False
    return True


def _color_label(c: Any) -> str:
    if not c:
        return "none"
    try:
        return "rgb(%d,%d,%d)" % tuple(int(round(x * 255)) for x in c)
    except Exception:
        return "none"


def _extract_geometry_bounded(page: Any, timeout: float):
    """(lines, curves) страницы с тайм-аутом; None — если страница «зависла»."""
    box: dict[str, Any] = {"done": False, "lines": [], "curves": []}

    def _job() -> None:
        try:
            box["lines"] = page.lines or []
            box["curves"] = page.curves or []
        except Exception:
            box["lines"], box["curves"] = [], []
        finally:
            box["done"] = True

    t = threading.Thread(target=_job, daemon=True)
    t.start()
    t.join(timeout)
    return (box["lines"], box["curves"]) if box["done"] else None


def _extract_plan_lengths(
    pdf: Any,
    page_texts: dict[int, str],
    slow_pages: set[int],
    timeout: float,
) -> tuple[list[dict[str, Any]], list[str]]:
    """Измеряет линии/кривые на страницах-планах PDF.

    Кабельные трассы выделяются по цвету (likely_cable=True). Длина — в единицах
    чертежа (pt); масштаб листа в PDF не хранится, поэтому перевод в метры
    выполняется только при явной надписи масштаба (см. check_plan_lengths).
    """
    lengths: list[dict[str, Any]] = []
    notes: list[str] = []
    for pn, ptext in page_texts.items():
        if pn in slow_pages:
            continue
        if not any(m in ptext.lower() for m in _PLAN_PAGE_MARKERS):
            continue
        if pn > len(pdf.pages):
            continue
        geo = _extract_geometry_bounded(pdf.pages[pn - 1], timeout)
        if geo is None:
            notes.append(f"Геометрия стр. {pn} не извлечена (тайм-аут).")
            continue
        lines, curves = geo
        for ln in lines:
            L = math.hypot(ln["x1"] - ln["x0"], ln["top"] - ln["bottom"])
            if L < 0.1:
                continue
            lengths.append(
                {
                    "layer": _color_label(ln.get("stroking_color")),
                    "type": "LINE",
                    "length": round(L, 3),
                    "likely_cable": _is_cable_color(ln.get("stroking_color")),
                }
            )
        for cu in curves:
            pts = cu.get("pts") or []
            L = 0.0
            for a, b in zip(pts, pts[1:]):
                L += math.hypot(b[0] - a[0], b[1] - a[1])
            if L < 0.1:
                continue
            lengths.append(
                {
                    "layer": _color_label(cu.get("stroking_color")),
                    "type": "CURVE",
                    "length": round(L, 3),
                    "likely_cable": _is_cable_color(cu.get("stroking_color")),
                }
            )
    return lengths, notes


# Только достаточно специфичные токены марок: короткие («ксп», «пвс», «нрг»…)
# срабатывают внутри обычных слов («эКСПлуатации») и дают ложные записи.
_JOURNAL_MARK_TOKENS = (
    "вббшв", "квббшв", "ввг", "кввг", "кгввг", "пугв", "шввп",
    "кспв", "кспэ", "ксвв", "ксбг", "ксбк",
    "тпп", "мкэш", "сип", "аввг",
    "frls", "frhf", "ftp", "utp", "nmf",
)


_SECTION_WORD_RE = re.compile(r"^\d+(?:[.,]\d+)?\s*[xх×*]\s*\d+(?:[.,]\d+)?$")


def _journal_line_to_item(line: str) -> dict[str, Any] | None:
    """Строка кабельного журнала из CAD-экспорта → запись (марка + сечение + длина).

    В PDF, экспортированных из CAD/XPS, таблица кабельного журнала не
    детектируется по колонкам, но текст строк содержит «… МАРКА СЕЧЕНИЕ ДЛИНА»:
    например «Е-1-7 XD1.1 XD1.2 11 2 ВВГнг(А)-LS 3х2,5 13» → марка ВВГнг(А)-LS,
    сечение 3х2,5, длина 13 м. Марка — целое слово с признаком кабельной
    продукции; сечение — следующее слово вида NxN; длина — последнее число.
    """
    t = re.sub(r"\s+", " ", line).strip()
    if len(t) < 8:
        return None
    nums = re.findall(r"\d+(?:[.,]\d+)?", t)
    if not nums:
        return None
    length = parse_float(nums[-1])
    if length is None:
        return None
    words = t.split()
    mark = None
    for w in words:
        wl = w.lower()
        if any(tok in wl for tok in _JOURNAL_MARK_TOKENS):
            mark = w
            break
    if not mark:
        return None
    sec_raw = ""
    try:
        mi = words.index(mark)
        for w in words[mi + 1 : mi + 3]:
            if _SECTION_WORD_RE.match(w):
                sec_raw = w.replace(",", ".").replace("х", "x").replace("×", "x").replace("*", "x")
                break
    except ValueError:
        pass
    section = None
    if sec_raw:
        parsed = parse_section(sec_raw)
        if parsed:
            section = parsed
    return {
        "name": mark,
        "mark": mark,
        "type": "",
        "section": section,
        "length": length,
        "qty": None,
        "from": "",
        "to": "",
        "sheet": "journal_text",
    }


def _words_to_lines(words: list, y_tol: float = 4.0) -> list[str]:
    """Группирует слова (x0,y0,x1,y1,word,...) в визуальные строки по y-координате.

    Нужно для кабельных журналов в CAD-экспортах: марка кабеля и длина находятся
    в одной визуальной строке, но разными текстовыми объектами, и get_text()
    выводит их на разных «строках» вывода.
    """
    words = sorted(words, key=lambda w: (w[1], w[0]))
    lines: list[list] = []
    cur: list = []
    cur_y: float | None = None
    for w in words:
        if cur_y is None or abs(w[1] - cur_y) <= y_tol:
            cur.append(w)
            if cur_y is None:
                cur_y = w[1]
        else:
            lines.append(cur)
            cur = [w]
            cur_y = w[1]
    if cur:
        lines.append(cur)
    out: list[str] = []
    for ln in lines:
        ln.sort(key=lambda w: w[0])
        out.append(" ".join(str(w[4]) for w in ln))
    return out


def _parse_journal_lines_pymupdf(path: Path, page_texts: dict[int, str]) -> tuple[list[dict[str, Any]], list[str]]:
    """Разбор кабельного журнала через координаты слов PyMuPDF.

    Для страниц с заголовком «Кабельный журнал» собирает визуальные строки
    (слова, выровненные по y) и распознаёт «… МАРКА … ДЛИНА» в каждой.
    """
    try:
        import pymupdf
    except ImportError:
        return [], ["PyMuPDF не установлен — журнал из текста не разобран."]
    journal: list[dict[str, Any]] = []
    notes: list[str] = []
    doc = pymupdf.open(str(path))
    try:
        for pn, ptext in page_texts.items():
            if "кабельный журнал" not in ptext.lower():
                continue
            if pn < 1 or pn > len(doc):
                continue
            words = doc[pn - 1].get_text("words")
            lines = _words_to_lines(words)
            hits = 0
            for ln in lines:
                it = _journal_line_to_item(ln)
                if it:
                    it["sheet"] = f"p{pn}_journal"
                    journal.append(it)
                    hits += 1
            if hits:
                notes.append(f"Кабельный журнал (стр. {pn}): {hits} строк распознано.")
    finally:
        doc.close()
    return journal, notes


def _parse_journal_lines(page_texts: dict[int, str]) -> tuple[list[dict[str, Any]], list[str]]:
    """Разбирает строки кабельного журнала из текста страниц с заголовком журнала."""
    journal: list[dict[str, Any]] = []
    notes: list[str] = []
    for pn, ptext in page_texts.items():
        low = ptext.lower()
        if "кабельный журнал" not in low:
            continue
        hits = 0
        for line in ptext.splitlines():
            it = _journal_line_to_item(line)
            if it:
                it["sheet"] = f"p{pn}_journal"
                journal.append(it)
                hits += 1
        if hits:
            notes.append(f"Кабельный журнал (стр. {pn}): {hits} строк распознано из текста.")
    return journal, notes


def _extract_tables_bounded(page: Any, timeout: float) -> list:
    """extract_tables() с жёстким тайм-аутом.

    pdfplumber умеет «зависать» на сложных страницах на десятки секунд.
    Возвращаем [] если не уложились; поток-демон дорабатывает в фоне и не
    блокирует ни сервер, ни последующие страницы.
    """
    box: dict[str, Any] = {"done": False, "tables": []}

    def _job() -> None:
        try:
            box["tables"] = page.extract_tables() or []
        except Exception:
            box["tables"] = []
        finally:
            box["done"] = True

    t = threading.Thread(target=_job, daemon=True)
    t.start()
    t.join(timeout)
    return list(box["tables"]) if box["done"] else []


def parse_pdf(path: Path) -> dict[str, Any]:
    from app.config import settings

    _quiet_pypdf()

    # 1) Текст — в отдельном процессе с жёстким бюджетом. Некоторые страницы
    #    (планы с плотной векторной графикой) разбираются pypdf десятки секунд;
    #    в потоке их не прервать, а утечки потоков выжигают бюджет разбора
    #    таблиц (спецификация/журнал оставались без таблиц). Процесс убивается
    #    целиком, поэтому таблицы спецификации разбираются стабильно.
    text_parts: list[str] = []
    tables: list[dict[str, Any]] = []
    notes = []
    slow_pages: set[int] = set()
    text_budget = float(getattr(settings, "pdf_text_budget", 60.0) or 60.0)
    text_page_timeout = float(getattr(settings, "pdf_text_page_timeout", 2.5) or 2.5)
    text_parts, t_notes = _extract_pdf_text_subprocess(path, text_budget, text_page_timeout)
    notes.extend(t_notes)

    # pypdf на PDF, экспортированных из CAD/XPS, может вернуть посимвольный
    # текст («И\nн\nв\n. ...») — тогда слова не склеиваются и спецификация/
    # журнал не распознаются. Переизвлекаем каскадом быстрых экстракторов:
    # PyMuPDF → pdftotext (poppler) → pdfplumber.
    if _text_is_garbled("\n".join(text_parts)) or not text_parts:
        if _text_is_garbled("\n".join(text_parts)):
            notes.append("Текст pypdf посимвольный (экспорт CAD/XPS) — переизвлечение.")
        mupdf_parts, t_notes = _extract_pymupdf_text_subprocess(path, text_budget)
        notes.extend(t_notes)
        if mupdf_parts:
            text_parts = mupdf_parts
        else:
            pdftotext_parts, t_notes = _extract_text_pdftotext(path, text_budget)
            notes.extend(t_notes)
            if pdftotext_parts:
                text_parts = pdftotext_parts
            else:
                text_parts, t_notes = _extract_pdfplumber_text_subprocess(path, text_budget, text_page_timeout)
                notes.extend(t_notes)

    if not text_parts:
        # запасной вариант: pdfplumber в потоке (если подпроцесс недоступен)
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages[:80]):
                    t = _extract_text_bounded(page, 3.0)
                    if t is None or not t.strip():
                        slow_pages.add(i + 1)
                        continue
                    text_parts.append(f"--- страница {i + 1} ---\n{t}")
        except Exception as exc:
            notes.append(f"pdfplumber: {exc}")

    # страницы, не попавшие в текст (медленные/векторные), помечаем —
    # их не даём и на разбор таблиц/геометрии
    have_text = set()
    for tp in text_parts:
        m = re.match(r"--- страница (\d+) ---\n?", tp)
        if m:
            have_text.add(int(m.group(1)))
    try:
        from pypdf import PdfReader

        total_pages = len(PdfReader(str(path)).pages)
    except Exception:
        total_pages = 80
    for pn in range(1, min(total_pages, 80) + 1):
        if pn not in have_text:
            slow_pages.add(pn)
    if slow_pages and not any("пропущен" in n for n in notes):
        notes.append(f"Текст {len(slow_pages)} стр. пропущен (сложная векторная графика).")

    text = "\n".join(text_parts).strip()

    # текст по страницам (нужен и для таблиц, и для геометрии планов)
    page_texts: dict[int, str] = {}
    for tp in text_parts:
        m = re.match(r"--- страница (\d+) ---\n?", tp)
        if m:
            page_texts[int(m.group(1))] = tp

    # 2) Таблицы — best-effort с жёсткими лимитами времени.
    #    Сначала обрабатываем страницы с явными табличными признаками
    #    (спецификация, кабельный журнал) — они быстрые; планы и схемы,
    #    где pdfplumber может «зависнуть», идут вторыми, если останется бюджет.
    if len(text) >= 40:
        try:
            import pdfplumber

            budget = float(getattr(settings, "pdf_table_budget", 15.0) or 15.0)
            page_timeout = float(getattr(settings, "pdf_table_page_timeout", 6.0) or 6.0)

            def _priority(pn: int) -> int:
                low = page_texts.get(pn, "").lower()
                # 0 — явные маркеры спецификации (ГОСТ 21.110) и кабельного журнала.
                #    «спецификац» ловит и обрыв строки «Спецификация оборудовани…»
                #    в CAD-экспорте. Эти страницы разбираются первыми.
                if any(
                    k in low
                    for k in (
                        "спецификац",
                        "кабельный журнал",
                        "направление кабеля",
                        "потребность кабелей",
                    )
                ):
                    return 0
                # 1 — таблицы вида «Поз. | Наименование | Кол. | Примечание»
                #    (перечни элементов, спецификации без явного заголовка)
                if "поз." in low and "наименование" in low:
                    return 1
                # 2 — прочие табличные страницы
                if "наименование" in low and "кол." in low:
                    return 2
                if any(k in low for k in ("наименование", "обозначение", "примечание", "марка", "длина")):
                    return 3
                return 4  # планы/чертежи, где «кол.» — из штампа: пропускаем

            candidates = sorted(
                [pn for pn in page_texts if _priority(pn) <= 3 and pn not in slow_pages],
                key=lambda pn: (_priority(pn), pn),
            )

            deadline = time.monotonic() + budget
            with pdfplumber.open(str(path)) as pdf:
                for pn in candidates:
                    if pn > len(pdf.pages):
                        continue
                    remaining = deadline - time.monotonic()
                    if remaining <= 0:
                        notes.append("Лимит времени на разбор таблиц PDF исчерпан — часть таблиц пропущена.")
                        break
                    page = pdf.pages[pn - 1]
                    extracted = _extract_tables_bounded(page, min(page_timeout, remaining))
                    for ti, table in enumerate(extracted):
                        if not table:
                            continue
                        rows = [[_cell(c) for c in row] for row in table if row]
                        if not rows:
                            continue
                        headers = rows[0]
                        mapped = _map_headers(headers)
                        records = []
                        for row in rows[1:]:
                            rec = {k: row[idx] if idx < len(row) else "" for k, idx in mapped.items()}
                            rec["_raw"] = row
                            records.append(rec)
                        tables.append(
                            {
                                "sheet": f"p{pn}_t{ti + 1}",
                                "headers": headers,
                                "mapped": mapped,
                                "records": records,
                            }
                        )
        except Exception as exc:
            notes.append(f"pdfplumber-tables: {exc}")

    if len(text) < 40:
        ocr_text, ocr_note = _ocr_pdf(path)
        notes.append(ocr_note)
        if ocr_text:
            text = ocr_text
        else:
            return _base(
                ok=False,
                kind="pdf",
                error="Из PDF не извлечён текст (вероятно сканированный чертёж). "
                + (ocr_note or "OCR недоступен."),
                notes="; ".join(notes),
            )

    items: list[dict[str, Any]] = []
    cables: list[dict[str, Any]] = []
    for table in tables:
        for rec in table.get("records", []):
            item = _record_to_item(rec, table.get("sheet", ""))
            if item:
                items.append(item)
                if looks_like_cable(_cable_blob(item)):
                    cables.append(item)

    # 3) Геометрия планов — измерение трасс (цветные линии/кривые).
    lengths: list[dict[str, Any]] = []
    if page_texts:
        try:
            import pdfplumber

            geo_timeout = float(getattr(settings, "pdf_text_page_timeout", 2.5) or 2.5) * 4
            with pdfplumber.open(str(path)) as pdf:
                lengths, geo_notes = _extract_plan_lengths(pdf, page_texts, slow_pages, geo_timeout)
                notes.extend(geo_notes)
        except Exception as exc:
            notes.append(f"pdf-geometry: {exc}")

    # 4) Кабельный журнал (CAD-экспорт): extract_tables не детектирует колонки
    #    журнала, но строки «… МАРКА … ДЛИНА» есть в тексте/координатах слов.
    journal_items: list[dict[str, Any]] = []
    journal_items, jn = _parse_journal_lines_pymupdf(path, page_texts)
    notes.extend(jn)
    if not journal_items:
        journal_items, jn = _parse_journal_lines(page_texts)
        notes.extend(jn)

    return _base(
        ok=True,
        kind="pdf",
        text=text[:MAX_TEXT],
        tables=tables,
        items=items,
        cables=cables,
        lengths=lengths,
        journal=journal_items,
        notes="; ".join(notes),
    )


def parse_dxf(path: Path) -> dict[str, Any]:
    try:
        import ezdxf
        from ezdxf import recover
    except ImportError as exc:
        return _base(ok=False, kind="dxf", error=f"ezdxf недоступен: {exc}")

    try:
        doc, _auditor = recover.readfile(str(path))
    except Exception:
        try:
            doc = ezdxf.readfile(str(path))
        except Exception as exc:
            return _base(ok=False, kind="dxf", error=f"Не удалось прочитать DXF: {exc}")

    return _from_dxf_doc(doc, kind="dxf")


def parse_dwg(path: Path) -> dict[str, Any]:
    # 1) ODA File Converter via ezdxf addon
    try:
        from ezdxf.addons import odafc

        doc = odafc.readfile(str(path))
        result = _from_dxf_doc(doc, kind="dwg")
        result["notes"] = "DWG преобразован через ODA File Converter."
        return result
    except Exception as exc:
        oda_err = str(exc)

    # 2) LibreDWG dwg2dxf
    dwg2dxf = shutil.which("dwg2dxf") or shutil.which("dwgread")
    if dwg2dxf:
        tmp = Path(tempfile.mkdtemp(prefix="dwg_"))
        try:
            out = tmp / (path.stem + ".dxf")
            cmd = [dwg2dxf, str(path), "-o", str(out)] if "dwg2dxf" in dwg2dxf else [dwg2dxf, "-O", "DXF", "-o", str(out), str(path)]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            if out.exists() and out.stat().st_size > 0:
                result = parse_dxf(out)
                result["kind"] = "dwg"
                result["notes"] = f"DWG преобразован через {Path(dwg2dxf).name}."
                return result
            lib_err = proc.stderr or proc.stdout or "пустой результат"
        except Exception as exc:
            lib_err = str(exc)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    else:
        lib_err = "dwg2dxf/dwgread не установлены"

    return _base(
        ok=False,
        kind="dwg",
        error=(
            "Формат DWG — закрытый. Чтение возможно только после преобразования в DXF. "
            f"ODA File Converter: {oda_err}. LibreDWG: {lib_err}. "
            "Установите ODA File Converter или сохраните чертежи в DXF."
        ),
    )


def _from_dxf_doc(doc, kind: str) -> dict[str, Any]:
    msp = doc.modelspace()
    texts: list[str] = []
    texts_geom: list[dict[str, Any]] = []
    equipment: list[dict[str, Any]] = []
    lengths: list[dict[str, Any]] = []

    def add_text(value: str, x: float = 0, y: float = 0, layer: str = ""):
        value = (value or "").strip()
        if not value:
            return
        texts.append(value)
        texts_geom.append({"text": value, "x": x, "y": y, "layer": layer})

    try:
        query_entities = list(msp)
    except Exception:
        query_entities = []

    for e in query_entities:
        dxftype = e.dxftype()
        layer = getattr(getattr(e, "dxf", None), "layer", "") or ""
        try:
            if dxftype == "TEXT":
                ins = e.dxf.insert
                add_text(e.dxf.text, float(ins[0]), float(ins[1]), layer)
            elif dxftype == "MTEXT":
                ins = e.dxf.insert
                add_text(e.text, float(ins[0]), float(ins[1]), layer)
            elif dxftype == "ATTRIB":
                ins = e.dxf.insert
                add_text(f"{e.dxf.tag}={e.dxf.text}", float(ins[0]), float(ins[1]), layer)
            elif dxftype == "INSERT":
                name = e.dxf.name
                ins = e.dxf.insert
                attrs = []
                try:
                    for attrib in e.attribs:
                        attrs.append(f"{attrib.dxf.tag}={attrib.dxf.text}")
                        add_text(attrib.dxf.text, float(ins[0]), float(ins[1]), layer)
                except Exception:
                    pass
                equipment.append(
                    {
                        "block": name,
                        "x": float(ins[0]),
                        "y": float(ins[1]),
                        "layer": layer,
                        "attrs": attrs,
                        "name": name,
                    }
                )
            elif dxftype in {"LWPOLYLINE", "POLYLINE", "LINE", "SPLINE"}:
                length = _entity_length(e)
                if length and length > 0.01:
                    lengths.append(
                        {
                            "layer": layer,
                            "type": dxftype,
                            "length": round(float(length), 3),
                            "likely_cable": _layer_is_cable(layer),
                        }
                    )
        except Exception:
            continue

    # paperspace texts
    try:
        for layout in doc.layouts:
            if layout.name.lower() == "model":
                continue
            for e in layout:
                if e.dxftype() in {"TEXT", "MTEXT"}:
                    try:
                        texts.append(e.dxf.text if e.dxftype() == "TEXT" else e.text)
                    except Exception:
                        pass
    except Exception:
        pass

    return _base(
        ok=True,
        kind=kind,
        text="\n".join(texts)[:MAX_TEXT],
        equipment=equipment,
        lengths=lengths,
        texts_geom=texts_geom[:5000],
        notes=f"Слоёв: {len(doc.layers) if hasattr(doc, 'layers') else '?'}; текстов: {len(texts)}; блоков: {len(equipment)}; линий: {len(lengths)}",
    )


def _entity_length(e) -> float | None:
    try:
        if e.dxftype() == "LINE":
            start, end = e.dxf.start, e.dxf.end
            dx = float(end[0]) - float(start[0])
            dy = float(end[1]) - float(start[1])
            dz = float(end[2]) - float(start[2]) if len(end) > 2 else 0.0
            return (dx * dx + dy * dy + dz * dz) ** 0.5
        if e.dxftype() == "LWPOLYLINE":
            return float(e.length())
        if e.dxftype() == "POLYLINE":
            if hasattr(e, "length"):
                return float(e.length())
            pts = list(e.points())
            total = 0.0
            for a, b in zip(pts, pts[1:]):
                total += ((a[0] - b[0]) ** 2 + (a[1] - b[1]) ** 2) ** 0.5
            return total
        if hasattr(e, "length"):
            return float(e.length())
    except Exception:
        return None
    return None


def _layer_is_cable(layer: str) -> bool:
    t = (layer or "").lower()
    keys = ("каб", "кабель", "cable", "трсс", "трасс", "эл ", "эом", "сс", "слабот", "пс", "соуэ", "скс")
    return any(k in t for k in keys)


def _ocr_pdf(path: Path) -> tuple[str, str]:
    from app.config import settings

    if not settings.ocr_enabled:
        return "", "OCR отключён в настройках."
    if not shutil.which("tesseract"):
        return "", "Tesseract OCR не установлен (пакет tesseract-ocr-rus)."
    try:
        from pdf2image import convert_from_path
    except ImportError:
        # fallback via pdftoppm
        if not shutil.which("pdftoppm"):
            return "", "Для OCR сканов нужны pdf2image или poppler (pdftoppm)."
        return _ocr_via_pdftoppm(path)
    try:
        import pytesseract
    except ImportError:
        pytesseract = None  # type: ignore
    try:
        pages = convert_from_path(str(path), dpi=150, first_page=1, last_page=8)
    except Exception as exc:
        return "", f"Растеризация PDF не удалась: {exc}"
    if pytesseract is None:
        return "", "Пакет pytesseract не установлен. OCR не выполнялся."
    parts = []
    for i, img in enumerate(pages):
        try:
            parts.append(f"--- OCR стр. {i + 1} ---\n" + pytesseract.image_to_string(img, lang="rus+eng"))
        except Exception as exc:
            parts.append(f"OCR стр. {i + 1}: {exc}")
    text = "\n".join(parts).strip()
    return text, "Текст получен OCR (возможны ошибки распознавания)." if text else "OCR не дал текста."


def _ocr_via_pdftoppm(path: Path) -> tuple[str, str]:
    tmp = Path(tempfile.mkdtemp(prefix="ocr_"))
    try:
        subprocess.run(
            ["pdftoppm", "-png", "-r", "150", "-f", "1", "-l", "8", str(path), str(tmp / "p")],
            check=False,
            timeout=90,
            capture_output=True,
        )
        parts = []
        for img in sorted(tmp.glob("p*.png")):
            proc = subprocess.run(
                ["tesseract", str(img), "stdout", "-l", "rus+eng"],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if proc.stdout.strip():
                parts.append(proc.stdout)
        text = "\n".join(parts).strip()
        return text, "OCR через pdftoppm+tesseract." if text else "OCR не дал текста."
    except Exception as exc:
        return "", f"OCR pdftoppm: {exc}"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _convert_with_libreoffice(path: Path, target: str) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    tmp = Path(tempfile.mkdtemp(prefix="lo_"))
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", target, "--outdir", str(tmp), str(path)],
            check=False,
            timeout=90,
            capture_output=True,
        )
        found = list(tmp.glob(f"*.{target}"))
        return found[0] if found else None
    except Exception:
        return None


HEADER_ALIASES: dict[str, tuple[str, ...]] = {
    "pos": ("поз", "№", "n", "nn", "п/п", "код", "position"),
    "name": ("наименование", "наименов", "название", "name", "оборудован", "наим"),
    "type": ("тип", "type", "исполнение"),
    "mark": ("марка", "тип, марка", "тип марка", "обознач", "марка кабеля", "кабель"),
    "section": ("сечен", "сечение", "жил", "мм2", "мм²", "q,", "qxn"),
    "unit": ("ед", "единиц", "изм", "unit"),
    "qty": ("кол", "колич", "qty", "количество", "число"),
    "length": ("длин", "l,", "трасс", "метр"),
    "from": ("откуда", "начало", "от куда", "start", "питающ"),
    "to": ("куда", "конец", "finish", "потребит", "назначен", "направлен"),
    "laying": ("способ", "проклад", "уклад", "лоток", "труба", "транше"),
    "power": ("мощн", "p,", "квт", "вт", "kw"),
    "current": ("ток", "iном", "а,", "current"),
    "voltage": ("напряж", "uном", "u,", "вольт", "voltage"),
    "cos": ("cos", "коэфф мощности", "cosφ", "cosf"),
    "note": ("примеч", "примечание", "note"),
    "manufacturer": ("завод", "изготов", "произв"),
}


def _header_row(rows: list[list[Any]]) -> int | None:
    best_i, best_s = None, 0
    for i, row in enumerate(rows[:12]):
        mapped = _map_headers(row)
        if len(mapped) > best_s:
            best_s, best_i = len(mapped), i
    return best_i if best_s >= 2 else None


def _map_headers(headers: list[Any]) -> dict[str, int]:
    mapped: dict[str, int] = {}
    for idx, raw in enumerate(headers):
        h = str(raw or "").strip().lower().replace("ё", "е")
        if not h:
            continue
        for key, aliases in HEADER_ALIASES.items():
            if key in mapped:
                continue
            # «Масса ед., кг» — это масса, а не единица измерения
            if key == "unit" and "масс" in h:
                continue
            if any(a in h for a in aliases):
                mapped[key] = idx
                break
    return mapped


def _cable_blob(item: dict[str, Any]) -> str:
    """Текст для определения «это кабель?» по всем значимым полям."""
    return " ".join(
        str(item.get(k) or "")
        for k in ("name", "mark", "type", "manufacturer", "note")
    )


def _record_to_item(rec: dict[str, Any], sheet: str) -> dict[str, Any] | None:
    name = str(rec.get("name") or "").strip()
    mark = str(rec.get("mark") or rec.get("type") or "").strip()
    if not name and not mark:
        raw = rec.get("_raw") or []
        joined = " ".join(str(x) for x in raw if x not in (None, ""))
        if len(joined) < 3:
            return None
        name = joined[:200]
    section = parse_section(" ".join(str(rec.get(k) or "") for k in ("section", "mark", "name", "type")))
    item = {
        "pos": str(rec.get("pos") or "").strip(),
        "name": name,
        "type": str(rec.get("type") or "").strip(),
        "mark": mark,
        "section": section,
        "unit": str(rec.get("unit") or "").strip(),
        "qty": parse_float(rec.get("qty")),
        "length": parse_float(rec.get("length")),
        "from": str(rec.get("from") or "").strip(),
        "to": str(rec.get("to") or "").strip(),
        "laying": str(rec.get("laying") or "").strip(),
        "power": parse_float(rec.get("power")),
        "current": parse_float(rec.get("current")),
        "voltage": parse_float(rec.get("voltage")),
        "cos": parse_float(rec.get("cos")),
        "note": str(rec.get("note") or "").strip(),
        "manufacturer": str(rec.get("manufacturer") or "").strip(),
        "sheet": sheet,
    }
    if item["length"] is None and item["unit"] in {"м", "м.", "м.п", "м.п.", "п.м", "пм"} and item["qty"]:
        item["length"] = item["qty"]
    return item


def _cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _base(**kwargs) -> dict[str, Any]:
    out = {
        "ok": False,
        "kind": "",
        "error": "",
        "notes": "",
        "text": "",
        "tables": [],
        "items": [],
        "cables": [],
        "equipment": [],
        "lengths": [],
        "texts_geom": [],
        "journal": [],
    }
    out.update(kwargs)
    if out.get("text"):
        out["text"] = truncate(out["text"], MAX_TEXT)
    return out
